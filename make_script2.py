from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.8)
    s.left_margin = s.right_margin = Cm(2.2)

BLACK = RGBColor(0x0D, 0x0D, 0x0D)
GREEN = RGBColor(0x00, 0x7A, 0x3D)
RED   = RGBColor(0xB8, 0x00, 0x00)
BLUE  = RGBColor(0x17, 0x4E, 0xA6)
GREY  = RGBColor(0x2C, 0x2C, 0x2C)
LGREY = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BGBLUE = "D9E8FB"
BGRED  = "FDE8E8"
BGGREY = "F4F4F4"


def set_para_bg(p, hex6):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    pPr.append(shd)


def set_cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)


def section(title, timing):
    """Bold green section header with timing."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "007A3D")
        pBdr.append(el)
    pPr.append(pBdr)
    set_para_bg(p, "E8F5EE")
    r1 = p.add_run(f"  {title.upper()}  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = GREEN
    r2 = p.add_run(f"  ({timing})")
    r2.font.size = Pt(9)
    r2.font.color.rgb = LGREY


def say(text):
    """What you say out loud — main script. Larger, dark, readable."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY


def cue(text):
    """On-screen action — blue background box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_para_bg(p, BGBLUE)
    r = p.add_run(f"  ▶  SHOW: {text}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE


def hit(text):
    """Pain line — slow down, let it land. Red background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    set_para_bg(p, BGRED)
    lq = "“"
    rq = "”"
    r = p.add_run(f"  {lq}{text}{rq}")
    r.italic = True
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RED


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(f"[ {text} ]")
    r.font.size = Pt(9)
    r.font.color.rgb = LGREY
    r.italic = True


def br():
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HUNTER")
r.bold = True; r.font.size = Pt(40); r.font.color.rgb = GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Demo Video Script  —  Full Version")
r.font.size = Pt(15); r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("hunter.dullugroup.co.ke  ·  Dullu Digital")
r.font.size = Pt(10); r.font.color.rgb = LGREY

br()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Runtime: 7–8 minutes  ·  Every feature covered  ·  No skips")
r.font.size = Pt(9); r.font.color.rgb = LGREY

doc.add_page_break()

# ── LEGEND ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("HOW TO READ THIS SCRIPT")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = GREEN

rows = [
    ("Normal text (large, dark)",     "Read this out loud. Speak naturally — not like a presenter, like a person."),
    ("▶  BLUE BOX",              "On-screen action. Click or show this before you say the next line."),
    ("RED ITALIC box",                "Key pain line. Pause before it. Say it slowly. Let it sit."),
    ("[ note ]",                      "Director note — do not read aloud."),
]
t = doc.add_table(rows=4, cols=2)
t.style = "Table Grid"
for i, (a, b) in enumerate(rows):
    c0, c1 = t.cell(i, 0), t.cell(i, 1)
    c0.text = a; c1.text = b
    for cell in (c0, c1):
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(c0, "E8F5EE")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 0 — OPENING / PAIN
# ══════════════════════════════════════════════════════════════════════════════
section("Opening — The Pain", "0:00 – 1:30")
note("Start recording. No product on screen yet. Just your voice and a blank or dark screen.")
br()

say("If you run a sales team, manage a book of clients, or you're out here selling anything to businesses — this is for you.")
br()
say("I want you to think about your last Monday morning.")
br()
say("How did your team start the week? Where did the first lead come from? How long did it take to find one business worth calling?")
br()

hit("Most sales teams spend more time finding people to talk to than actually talking to them.")

br()
say("Every morning, your team opens Google Maps. Types in a business type. Scrolls. Copies a number. Goes back. Does it again.")
br()

hit("Four hours later, you haven't called anyone.")

br()
say("And when you do finally get through — you're calling with a script that sounds exactly like the last three people who called that number.")
br()
say("No context. No research. Just: 'Hi, I'm calling from XYZ, we offer…'")
br()

hit("And the person on the other end? They've already mentally hung up before you finish the sentence.")

br()
say("This is the problem whether you're selling insurance, running a digital agency, working in a bank selling SME products, or pushing solar to businesses.")
br()
say("The leads are out there. The market is real. The money is there.")
br()
say("But between you and that money is a wall: finding the right person, knowing enough about their business to say something worth listening to, and then actually following up before someone else does.")
br()

hit("That wall is what kills pipelines. And most teams have just accepted it as part of the job.")

br()
say("It doesn't have to be. Let me show you what we built.")
br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — WHAT IS HUNTER
# ══════════════════════════════════════════════════════════════════════════════
section("Part 1 — What Is Hunter", "1:30 – 2:00")

cue("Open browser — navigate to hunter.dullugroup.co.ke. Show the dashboard.")
br()

say("This is Hunter.")
br()
say("Hunter is an AI-powered lead intelligence and outreach engine built for B2B sales teams.")
br()
say("It finds businesses, researches them automatically, scores them by how likely they are to buy, writes a personalised opening message, and puts everything in a pipeline your team can manage.")
br()
say("What used to take a salesperson three to four hours every morning — Hunter does in under a minute.")
br()

hit("You don't need a bigger team. You need your current team working smarter.")

br()
say("Let me walk you through exactly how it works — from zero to a ready-to-send message.")
br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — LEAD DISCOVERY
# ══════════════════════════════════════════════════════════════════════════════
section("Part 2 — Finding Leads", "2:00 – 3:00")

cue("Go to the Search / Scrape section. Show the input fields.")
br()

say("Step one: tell Hunter what kind of business you're looking for and where.")
br()
say("Let's say I'm selling group medical insurance to companies in Nairobi. I type in 'clinic' or 'hospital' and select Nairobi.")
br()

cue("Type a vertical (e.g. 'dental clinic') and select a city. Hit Search.")
br()

say("Hunter goes out to Google Places, finds real registered businesses, and pulls everything you need: business name, phone number, website, physical address, their Google rating, and how many reviews they have.")
br()

cue("Show the lead cards loading in — wait for results to populate.")
br()

say("No spreadsheets. No copy-paste. No scrolling through Google Maps for an hour.")
br()
say("These are real, verified businesses. Each one with a phone number you can call right now.")
br()

hit("Your competitor's rep is still building their list. Yours is already dialling.")

br()

cue("Scroll through the lead cards slowly. Show the rating, review count, phone, and website visible on each card.")
br()

say("And notice — before we've even opened a single business, we already know which ones are established, which ones have high customer volume, and which ones are worth prioritising.")
br()
say("That's the Google rating and review count doing work before we even pick up the phone.")
br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 — ENRICHMENT MODES
# ══════════════════════════════════════════════════════════════════════════════
section("Part 3 — Enrichment Modes (What You're Selling Matters)", "3:00 – 3:45")

cue("Open the Enrichment Mode selector. Scroll slowly through all 8 modes.")
br()

say("Here's where Hunter is different from any other lead tool.")
br()
say("Before it researches a business, you tell it what you're selling. Because what matters for a digital agency is completely different from what matters for an insurance broker or a bank.")
br()
say("Hunter has eight modes built in:")
br()
say("Digital and agency sales. Insurance. Fintech and lending. Solar and clean energy. Telecom and connectivity. Healthcare and pharma. Recruitment and staffing. And general B2B.")
br()
say("Each mode tells Hunter what signals to look for, what gaps to flag, and what angle your message should take. It's already calibrated to your buyer, not a generic template.")
br()

hit("Every other tool gives you a name and a number. Hunter gives you a briefing.")

br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 — BUSINESS INTELLIGENCE / ENRICHMENT
# ══════════════════════════════════════════════════════════════════════════════
section("Part 4 — Business Intelligence", "3:45 – 5:00")

cue("Click into a single lead. Show the full lead detail page.")
br()

say("Now watch what happens when we open one business.")
br()
say("Hunter has already visited their website and read it. It checked their tech stack. It looked at whether they have a booking system, live chat, online payments, SSL, a WhatsApp number, a Meta pixel.")
br()

cue("Scroll down to the Business Intelligence panel. Show the signals detected — the green ticks and red gaps.")
br()

say("These green signals are things they have. These red gaps are things they're missing.")
br()
say("If I'm selling digital services — every red gap is a door. No booking system. No online payments. No live chat. That's money leaving their business every single day.")
br()
say("If I'm selling insurance — Hunter looks at how many staff they seem to have, whether they run deliveries, whether they have a physical premises. Everything that tells me what cover they need before I say a word.")
br()
say("If I'm in fintech, selling business loans — Hunter looks for growth signals. New branches. High review volume. Expansion language. These are businesses that need capital and can repay it.")
br()

hit("You walk into every call knowing what their business looks like from the outside. That changes the entire conversation.")

br()

cue("Show the AI Score on the lead — the number out of 100.")
br()

say("And Hunter scores each lead out of a hundred. That score tells you: how much pain is this business in, and how much of that pain can you solve?")
br()
say("A score above seventy means this business has real gaps and real budget. That's your priority call.")
br()
say("A score below forty means there's not much to work with right now. Move on. Don't waste the call.")
br()

hit("Stop wasting your best salesperson on leads that were never going to convert.")

br()

cue("Show the contact extraction section — the decision-maker name and title if found.")
br()

say("And if Hunter found a decision-maker's name or title from their website or listings — it shows you that too. So you're not calling and asking 'who handles this' — you already know who to ask for.")
br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 5 — AI OUTREACH
# ══════════════════════════════════════════════════════════════════════════════
section("Part 5 — The Outreach Message", "5:00 – 6:00")

cue("Click 'Generate Outreach' on the lead. Let the message stream in — don't skip this, let them watch it write itself.")
br()

say("Now this is the part that changes everything about cold outreach.")
br()
say("Watch this message.")
br()
say("It's not a template. It's not 'Hi, I hope this finds you well.' Hunter has read their website, identified their specific gap, and written an opening that references exactly what that business is missing and what it's costing them.")
br()

cue("Pause. Read the generated message out loud slowly.")
br()

say("Listen to that. It mentions their business by name. It references a specific problem they have. And it makes the cost of that problem concrete.")
br()

hit("That message gets read. A template gets deleted.")

br()
say("The reason generic outreach doesn't work is that businesses are drowning in it. Everyone sounds the same. Hunter makes you sound like you actually did your homework — because it did.")
br()

cue("Toggle between WhatsApp and Email versions of the message.")
br()

say("You can send it as a WhatsApp message or as an email. Both versions are written. You pick the channel, copy, and send.")
br()
say("One click. Message ready. Specific to that business. In the right format for the right channel.")
br()

hit("Your team goes from spending an hour writing one outreach to sending fifty personalised messages in the same time.")

br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 6 — CALL GUIDE
# ══════════════════════════════════════════════════════════════════════════════
section("Part 6 — The Call Guide", "6:00 – 6:50")

cue("Open the Call Guide tab on the same lead.")
br()

say("Let's say they responded. Or you're calling cold. Either way, your salesperson now opens the Call Guide.")
br()
say("This is a live script — built around what Hunter found about that specific business. It walks your rep through the call stage by stage.")
br()

cue("Show the call stages: Opening → Qualify → Pitch → Handle Objection → Close.")
br()

say("Opening: how to introduce yourself without sounding like a cold caller.")
br()
say("Qualify: two or three questions to confirm this is worth continuing.")
br()
say("Pitch: what to say based on the specific gap Hunter found. Not a generic product pitch — a specific problem-solution.")
br()
say("And if they object — the guide shows exactly how to respond to the most common pushbacks. 'We're not interested.' 'Send me something.' 'We already have someone.' All handled.")
br()

cue("Click through an objection handling node. Show the branching logic.")
br()

say("This is not a rigid script. It branches based on how the prospect responds. Your rep follows the conversation, not a monologue.")
br()

hit("A junior salesperson with three months' experience walks into that call with the confidence of someone who's done it a hundred times.")

br()
say("That's what training on a good tool does. And that's what Hunter provides on every single lead.")
br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 7 — PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
section("Part 7 — The Pipeline", "6:50 – 7:20")

cue("Go back to the leads list. Show the stage selector on a lead card.")
br()

say("Every lead Hunter finds sits in your pipeline. And every lead has a stage.")
br()
say("Called. Qualified. Pitched. Booked. Closed. Or not interested — so you know not to waste another call.")
br()

cue("Move a lead from 'Called' to 'Qualified'. Show it updating.")
br()

say("Your manager can see this in real time. Where are deals moving? Where are they stalling? Where is the team losing business and why?")
br()
say("No more tracking deals in a WhatsApp group. No more 'did anyone follow up with that hotel last week?'")
br()

hit("Leads that fall through the cracks are revenue you lost and don't know you lost.")

br()
say("Hunter makes sure nothing falls through. Every lead has a status. Every salesperson knows what they're working and what's next.")
br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 8 — SPECIFIC PAIN BY INDUSTRY
# ══════════════════════════════════════════════════════════════════════════════
section("Part 8 — Specific to Your Industry", "7:20 – 7:50")
note("This section is optional based on your audience. Mention two or three that match who's watching.")
br()

say("Let me be specific for a moment, because Hunter is built around real industry pain — not generic sales advice.")
br()
say("If you're in insurance: your best reps spend sixty to seventy percent of their week on cold prospecting. Hunter does that part. They spend their week closing.")
br()
say("If you're in banking, selling SME products: branch-led banks are losing SME clients to digital lenders on speed. Hunter gets you into the conversation first — before the digital lender does.")
br()
say("If you run a digital agency: you win clients at the moment they decide to spend. Hunter puts you in their inbox at that exact moment, with a message that shows you already understand their gaps.")
br()
say("If you're in recruitment: the companies that are growing are hiring. Hunter finds them by looking for expansion signals, new locations, high volume activity. You reach them before they've posted the job.")
br()
say("If you're in solar or telecom: Hunter finds the businesses with the highest energy or connectivity pain — 24-hour operations, large premises, generator dependency, no SSL on their site. The ones already feeling the problem.")
br()

hit("Every vertical has a version of this problem. Hunter is built to solve it for yours.")

br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CLOSE
# ══════════════════════════════════════════════════════════════════════════════
section("Close", "7:50 – 8:30")

cue("Go back to the dashboard. Show a full view of the lead pipeline — multiple scored, enriched leads.")
br()

say("This is what your team's morning looks like with Hunter.")
br()
say("Not an hour on Google Maps. Not generic messages that get ignored. Not calling blind.")
br()
say("Researched leads. Scored by priority. With a personalised message ready to send and a call guide ready to open.")
br()

hit("You're not getting more hours in the day. But you are getting more out of the ones you have.")

br()
say("Hunter is live. It's running right now.")
br()
say("If you want to see it pull real leads from your industry, your city, with your vertical — send me a message. I'll show you a live demo with your actual market.")
br()
say("Not a pitch deck. Not a presentation. Real leads. Real research. Your businesses.")
br()

hit("The leads are there. The question is whether you reach them first.")

br()
note("Pause two seconds. End recording.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run("QUICK REFERENCE — SCREEN CUES")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GREEN

rows2 = [
    ("Time",       "What you say",                              "What to show on screen"),
    ("0:00–1:30",  "Pain — universal prospecting problem",    "Blank screen or face. No product yet."),
    ("1:30–2:00",  "What Hunter is",                            "Hunter dashboard homepage"),
    ("2:00–3:00",  "Finding leads",                             "Search input → leads loading → lead cards with phone/rating/website"),
    ("3:00–3:45",  "Enrichment modes",                          "Mode selector — scroll all 8 slowly"),
    ("3:45–5:00",  "Business intelligence",                     "Single lead open → signals panel → score → contact extraction"),
    ("5:00–6:00",  "AI outreach message",                       "Generate Outreach → message streaming → WhatsApp/Email toggle"),
    ("6:00–6:50",  "Call guide",                                "Call Guide tab → stage flow → objection node"),
    ("6:50–7:20",  "Pipeline",                                  "Lead list → stage selector → move a lead through stages"),
    ("7:20–7:50",  "Industry-specific pain",                    "Mode selector again or stay on dashboard"),
    ("7:50–8:30",  "Close + CTA",                               "Full dashboard — multiple enriched leads visible"),
]

t2 = doc.add_table(rows=len(rows2), cols=3)
t2.style = "Table Grid"
for i, row in enumerate(rows2):
    for j, txt in enumerate(row):
        cell = t2.cell(i, j)
        cell.text = txt
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        if i == 0:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_bg(cell, "007A3D")
            cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        else:
            set_cell_bg(cell, "F2F2F2" if i % 2 == 0 else "FFFFFF")

br()
p = doc.add_paragraph()
r = p.add_run("KEY PAIN LINES — SLOW DOWN ON THESE")
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RED

pains = [
    "Most sales teams spend more time finding people to talk to than actually talking to them.",
    "That wall is what kills pipelines. And most teams have just accepted it as part of the job.",
    "You walk into every call knowing what their business looks like from the outside. That changes the entire conversation.",
    "Stop wasting your best salesperson on leads that were never going to convert.",
    "That message gets read. A template gets deleted.",
    "Your team goes from spending an hour writing one outreach to sending fifty personalised messages in the same time.",
    "A junior salesperson with three months' experience walks into that call with the confidence of someone who's done it a hundred times.",
    "Leads that fall through the cracks are revenue you lost and don't know you lost.",
    "Every vertical has a version of this problem. Hunter is built to solve it for yours.",
    "The leads are there. The question is whether you reach them first.",
]
for pl in pains:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    set_para_bg(p, BGRED)
    lq = "“"
    rq = "”"
    r = p.add_run(f"  {lq}{pl}{rq}")
    r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = RED

br()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("hunter.dullugroup.co.ke  ·  @iandullu  ·  Dullu Digital")
r.font.size = Pt(9); r.font.color.rgb = LGREY

out = "/home/dullz/hunter-saas/Hunter_Demo_Script.docx"
doc.save(out)
print(f"Saved: {out}")
